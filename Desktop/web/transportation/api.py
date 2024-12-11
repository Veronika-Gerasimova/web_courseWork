from rest_framework.viewsets import GenericViewSet
from django.contrib.auth.models import User
from rest_framework import mixins
from transportation.models import Client, Flight, Ticket, Baggage, Airplane, UserProfile
from transportation.serializers import ClientSerializer, FlightSerializer, TicketSerializer, BaggageSerializer, AirplaneSerializer,  UserLoginSerializer
from rest_framework.decorators import action
from rest_framework.response import Response
from django.db.models import Count, Avg, Max, Min
from rest_framework.permissions import IsAuthenticated, BasePermission
from django.core.cache import cache
from django.contrib.auth import authenticate, login
import pyotp
from rest_framework import serializers, viewsets, status
from django.contrib.auth import authenticate, login, logout
from django.views.decorators.csrf import csrf_exempt 
import os
from django.http import FileResponse
from openpyxl import Workbook  # для Excel
from docx import Document  # для Word
from rest_framework.permissions import BasePermission
from django.core.cache import cache
from django.http import HttpResponse
import qrcode
from io import BytesIO
from rest_framework.viewsets import ModelViewSet
# Миксин для фильтрации данных по пользователю
class UserFilteredViewSet(GenericViewSet):
    def get_queryset(self):
        user = self.request.user
        if user.is_superuser:
            # Для суперпользователя показываем все
            return super().get_queryset()
        else:
            # Для обычного пользователя фильтруем по его user_id
            return self.queryset.filter(user=user)

class ClientViewSet(
    mixins.CreateModelMixin,
    mixins.UpdateModelMixin,
    mixins.RetrieveModelMixin,
    mixins.ListModelMixin,
    mixins.DestroyModelMixin,
    UserFilteredViewSet
):
    queryset = Client.objects.all()
    serializer_class = ClientSerializer
    def update(self, request, *args, **kwargs):
        try:
            instance = self.get_object()
            serializer = self.get_serializer(instance, data=request.data, partial=True)
            serializer.is_valid(raise_exception=True)
            self.perform_update(serializer)
            return Response(serializer.data)
        except Exception as e:
            return Response({'detail': str(e)}, status=500)
    @action(detail=False, methods=["GET"], url_path="stats")
    def get_stats(self, request, *args, **kwargs):
        stats = Client.objects.aggregate(
            count=Count("id"),
            avg=Avg("id"),
            max=Max("id"),
            min=Min("id"),
        )
        return Response(stats)

    @action(detail=False, methods=["GET"], url_path="export")
    def export_clients(self, request, *args, **kwargs):
        file_type = request.query_params.get("type", "excel")  # Тип файла: excel или word

        if file_type == "word":
            document = Document()
            document.add_heading('Список клиентов', level=1)

            for client in Client.objects.all():
                document.add_paragraph(f"ФИО: {client.name}")
                document.add_paragraph(f"Email: {client.email}")
                document.add_paragraph(f"Телефон: {client.phone}")
                document.add_paragraph("-" * 40)

            file_path = "clients.docx"
            document.save(file_path)

        else:  # По умолчанию Excel
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = "Клиенты"
            sheet.append(["ФИО", "Email", "Телефон"])

            for client in Client.objects.all():
                sheet.append([client.name, client.email, client.phone])

            file_path = "clients.xlsx"
            workbook.save(file_path)

        response = FileResponse(open(file_path, 'rb'), as_attachment=True)
        response['Content-Disposition'] = f'attachment; filename="{os.path.basename(file_path)}"'

        return response
        
class FlightViewSet(
    mixins.CreateModelMixin,
    mixins.UpdateModelMixin,
    mixins.RetrieveModelMixin,
    mixins.ListModelMixin,
    mixins.DestroyModelMixin,
    UserFilteredViewSet
):
    queryset = Flight.objects.all()
    serializer_class = FlightSerializer

    @action(detail=False, methods=["GET"], url_path="stats")
    def get_stats(self, request, *args, **kwargs):
        stats = Flight.objects.aggregate(
            count=Count("id"),
            avg=Avg("id"),
            max=Max("id"),
            min=Min("id"),
        )
        return Response(stats)


class TicketViewSet(
    mixins.CreateModelMixin,
    mixins.UpdateModelMixin,
    mixins.RetrieveModelMixin,
    mixins.ListModelMixin,
    mixins.DestroyModelMixin,
    UserFilteredViewSet
):
    queryset = Ticket.objects.all()
    serializer_class = TicketSerializer

    @action(detail=False, methods=["GET"], url_path="stats")
    def get_stats(self, request, *args, **kwargs):
        stats = Ticket.objects.aggregate(
            count=Count("id"),
            avg=Avg("id"),
            max=Max("id"),
            min=Min("id"),
        )
        return Response(stats)


class BaggageViewSet(
    mixins.CreateModelMixin,
    mixins.UpdateModelMixin,
    mixins.RetrieveModelMixin,
    mixins.ListModelMixin,
    mixins.DestroyModelMixin,
    UserFilteredViewSet
):
    queryset = Baggage.objects.all()
    serializer_class = BaggageSerializer

    @action(detail=False, methods=["GET"], url_path="stats")
    def get_stats(self, request, *args, **kwargs):
        stats = Baggage.objects.aggregate(
            count=Count("id"),
            avg=Avg("id"),
            max=Max("id"),
            min=Min("id"),
        )
        return Response(stats)


class AirplaneViewSet(
    mixins.CreateModelMixin,
    mixins.UpdateModelMixin,
    mixins.RetrieveModelMixin,
    mixins.ListModelMixin,
    mixins.DestroyModelMixin,
    UserFilteredViewSet
):
    queryset = Airplane.objects.all()
    serializer_class = AirplaneSerializer

    @action(detail=False, methods=["GET"], url_path="stats")
    def get_stats(self, request, *args, **kwargs):
        stats = Airplane.objects.aggregate(
            count=Count("id"),
            avg=Avg("id"),
            max=Max("id"),
            min=Min("id"),
        )
        return Response(stats)
  
class UserViewSet(viewsets.GenericViewSet):
    serializer_class = UserLoginSerializer
    permission_classes = [IsAuthenticated]

    @action(url_path="info", methods=["POST"], detail=False)
    def get_info(self, request, *args, **kwargs):
        user = request.user
        if user.is_authenticated:
            data = {
                'is_authenticated': True,
                'username': user.username,
                'user_id': user.id,
            }
            return Response(data)
        return Response({'detail': 'Unauthorized'}, status=401)

    @action(url_path="login", methods=["POST"], detail=False)
    def login(self, request, *args, **kwargs):
        serializer = self.get_serializer(data=request.data)
        if not serializer.is_valid():
            return Response(serializer.errors, status=400)  # Возврат ошибок валидации
        
        user = authenticate(request, username=serializer.validated_data['username'], password=serializer.validated_data['password'])
        
        if user:
            # Создаем профиль пользователя, если его нет
            if not hasattr(user, 'profile'):
                user.profile = UserProfile.objects.create(user=user)
            
            login(request, user)
            return Response({"detail": "Login successful"})
        
        return Response({"detail": "Invalid credentials"}, status=401)

    @action(url_path="logout", methods=["POST"], detail=False)
    def logout(self, request, *args, **kwargs):
        logout(request)
        return Response({"detail": "Logout successful"})

    @action(url_path="register", methods=["POST"], detail=False)
    def register(self, request, *args, **kwargs):
        username = request.data.get('username')
        password = request.data.get('password')
        if username and password:
            user = User.objects.create_user(username=username, password=password)
            return Response({"detail": "Registration successful"})
        return Response({"detail": "Username and password are required"}, status=400)


def generate_otp_key(user):
    # Проверяем, есть ли уже ключ у пользователя
    if not user.profile.otp_key:
        # Генерируем новый OTP ключ для пользователя
        totp = pyotp.TOTP(pyotp.random_base32())  # Генерация нового случайного ключа
        user.profile.otp_key = totp.secret  # Сохраняем его в профиле пользователя
        user.profile.save()  # Сохраняем изменения


class SecuredModelViewSet(ModelViewSet):
    class OTPSerializer(serializers.Serializer):
        key = serializers.CharField()

    class OTPRequired(BasePermission):
        """Проверка действующего OTP-токена"""
        def has_permission(self, request, view):
            return bool(cache.get(f'otp_good_{request.user.id}', False))

    @action(detail=False, methods=['POST'], url_path='otp-login', serializer_class=OTPSerializer)
    def otp_login(self, request, *args, **kwargs):
        generate_otp_key(request.user)
        totp = pyotp.TOTP(request.user.profile.otp_key)  # `otp_key` должен быть у пользователя
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)

        success = False
        if totp.verify(serializer.validated_data['key']):
            cache.set(f'otp_good_{request.user.id}', True, timeout=300)  # Время жизни OTP (300 сек.)
            success = True

        return Response({'success': success}, status=status.HTTP_200_OK)

    @action(detail=False, methods=['GET'], url_path='otp-status')
    def otp_status(self, request, *args, **kwargs):
        otp_good = cache.get(f'otp_good_{request.user.id}', False)
        return Response({'otp_good': otp_good}, status=status.HTTP_200_OK)

    def get_permissions(self):
        """Добавляем проверку на OTP для редактирования."""
        if self.action in ['update', 'partial_update', 'destroy']:
            self.permission_classes = [IsAuthenticated, self.OTPRequired]
        return super().get_permissions()
    
    @action(detail=False, methods=['GET'], url_path='otp-qr-code')
    def generate_qr_code(self, request, *args, **kwargs):
        # Убедитесь, что пользователь имеет профиль с OTP-ключом
        if not hasattr(request.user, 'profile') or not request.user.profile.otp_key:
            print("Ошибка: профиль или ключ OTP отсутствуют.")
            return Response({"error": "Профиль пользователя не настроен."}, status=400)


        # Создание OTP URI
        totp = pyotp.TOTP(request.user.profile.otp_key)
        otp_uri = totp.provisioning_uri(
            name=request.user.username,
            issuer_name="Travel"
        )

        # Создание QR-кода
        qr = qrcode.make(otp_uri)
        buffer = BytesIO()
        qr.save(buffer, format="PNG")
        buffer.seek(0)

        # Возврат QR-кода в формате изображения
        return HttpResponse(buffer, content_type="image/png")
